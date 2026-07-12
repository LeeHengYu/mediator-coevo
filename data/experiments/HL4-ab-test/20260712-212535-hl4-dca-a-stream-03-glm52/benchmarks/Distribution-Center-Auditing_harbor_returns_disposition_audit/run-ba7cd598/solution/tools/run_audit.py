#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
PLAN_XLSX = ROOT_DIR / 'Return_Plan.xlsx'
EVENT_XLSX = ROOT_DIR / 'Disposition_Event_Log.xlsx'
ALIAS_XLSX = ROOT_DIR / 'Disposition_Alias.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Returns_Disposition_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Returns_Disposition_Brief.docx'

BASE_HEADERS = [
    'Return ID',
    'Line ID',
    'Planned Disposition',
    'Reason Code',
    'Requested Qty',
    'Warehouse',
    'Carrier',
    'Lane',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Missing Final Event', 'Disposition Mismatch', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Warehouse', 'Carrier', 'Missing Final Events', 'Disposition Mismatches', 'Total Errors']


def normalize_disposition(raw, alias_map):
    key = str(raw or '').strip().upper()
    return alias_map.get(key, key)


def calc_flags(row, latest_completed, alias_map):
    key = (str(row['Return ID']), str(row['Line ID']))
    event = latest_completed.get(key)
    missing = 1 if event is None else 0
    planned = str(row['Planned Disposition'] or '').strip().upper()
    mismatch = 0
    if event is not None:
        normalized = normalize_disposition(event[1], alias_map)
        mismatch = 1 if normalized != planned else 0
    total = missing + mismatch
    if total == 0:
        summary = 'None'
    elif missing and mismatch:
        summary = 'Missing Final Event, Disposition Mismatch'
    elif missing:
        summary = 'Missing Final Event'
    else:
        summary = 'Disposition Mismatch'
    return missing, mismatch, total, summary


plan_wb = load_workbook(PLAN_XLSX, data_only=True)
event_wb = load_workbook(EVENT_XLSX, data_only=True)
alias_wb = load_workbook(ALIAS_XLSX, data_only=True)

plan_ws = plan_wb['PlanLines']
event_ws = event_wb['Events']
alias_ws = alias_wb['AliasMap']

headers = [plan_ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, plan_ws.max_row + 1):
    values = [plan_ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

latest_completed = {}
for r in range(2, event_ws.max_row + 1):
    return_id = event_ws.cell(r, 1).value
    line_id = event_ws.cell(r, 2).value
    event_time = event_ws.cell(r, 3).value
    status = str(event_ws.cell(r, 5).value or '').strip().upper()
    final_disp = event_ws.cell(r, 6).value
    if return_id is None or line_id is None or status != 'COMPLETED':
        continue
    key = (str(return_id), str(line_id))
    record = (str(event_time), str(final_disp))
    if key not in latest_completed or record[0] > latest_completed[key][0]:
        latest_completed[key] = record

alias_map = {}
for r in range(2, alias_ws.max_row + 1):
    alias = alias_ws.cell(r, 1).value
    standard = alias_ws.cell(r, 2).value
    if alias is None or standard is None:
        continue
    alias_map[str(alias).strip().upper()] = str(standard).strip().upper()

out_wb = Workbook()
raw_ws = out_wb.active
raw_ws.title = 'RawData'
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
return_totals = defaultdict(int)
total_missing = total_mismatch = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    missing, mismatch, total, summary = calc_flags(row, latest_completed, alias_map)
    formatted_ws.append(base_values + [missing, mismatch, total, summary])
    key = (str(row['Warehouse']), str(row['Carrier']))
    agg[key][0] += missing
    agg[key][1] += mismatch
    agg[key][2] += total
    return_totals[str(row['Return ID'])] += total
    total_missing += missing
    total_mismatch += mismatch
    total_errors += total

for (warehouse, carrier), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([warehouse, carrier, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_missing, total_mismatch, total_errors])
out_wb.save(OUTPUT_XLSX)

top_returns = [rid for rid, total in sorted(return_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('Missing Final Event checks whether a return line has no completed disposition event. Disposition Mismatch checks whether the final disposition does not match the planned disposition after alias normalization.')
doc.add_paragraph(f'The audit found {total_missing} Missing Final Events, {total_mismatch} Disposition Mismatches, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority return IDs include {", ".join(top_returns)}. Recommendation: enforce event completion workflows and validate disposition mapping tables.')
doc.save(OUTPUT_DOCX)
