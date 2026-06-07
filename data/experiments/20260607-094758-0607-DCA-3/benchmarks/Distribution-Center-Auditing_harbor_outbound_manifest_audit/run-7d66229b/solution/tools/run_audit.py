#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
PLAN_XLSX = ROOT_DIR / 'Manifest_Plan.xlsx'
SCAN_XLSX = ROOT_DIR / 'Dock_Scan_Log.xlsx'
TEMPLATE_XLSX = ROOT_DIR / 'Outbound_Audit_Template.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Outbound_Load_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Outbound_Load_Brief.docx'

BASE_HEADERS = [
    'Shipment ID',
    'Carton ID',
    'Planned Zone',
    'Route',
    'Expected Weight',
    'Hazmat Flag',
    'Carrier',
    'Wave',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Missing Load Scan', 'Zone Mismatch', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Route', 'Shipment ID', 'Missing Load Scans', 'Zone Mismatches', 'Total Errors']

def calc_flags(row, latest_loaded):
    key = (str(row['Shipment ID']), str(row['Carton ID']))
    scan = latest_loaded.get(key)
    missing_scan = 1 if scan is None else 0
    zone_mismatch = 0 if scan is None else int(str(scan[1]) != str(row['Planned Zone']))
    total = missing_scan + zone_mismatch
    if total == 0:
        summary = 'None'
    elif missing_scan and zone_mismatch:
        summary = 'Missing Load Scan, Zone Mismatch'
    elif missing_scan:
        summary = 'Missing Load Scan'
    else:
        summary = 'Zone Mismatch'
    return missing_scan, zone_mismatch, total, summary

plan_wb = load_workbook(PLAN_XLSX, data_only=True)
scan_wb = load_workbook(SCAN_XLSX, data_only=True)
out_wb = load_workbook(TEMPLATE_XLSX)

plan_ws = plan_wb['PlanLines']
headers = [plan_ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, plan_ws.max_row + 1):
    values = [plan_ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

latest_loaded = {}
scan_ws = scan_wb['Scans']
for r in range(2, scan_ws.max_row + 1):
    shipment_id = scan_ws.cell(r, 1).value
    carton_id = scan_ws.cell(r, 2).value
    scanned_zone = scan_ws.cell(r, 3).value
    scan_ts = scan_ws.cell(r, 4).value
    status = str(scan_ws.cell(r, 5).value or '').strip().upper()
    if shipment_id is None or carton_id is None or status != 'LOADED':
        continue
    key = (str(shipment_id), str(carton_id))
    record = (str(scan_ts), str(scanned_zone))
    if key not in latest_loaded or record[0] > latest_loaded[key][0]:
        latest_loaded[key] = record

for sheet_name, header_row in [('RawData', BASE_HEADERS), ('Formatted Data', FORMATTED_HEADERS), ('Summary', SUMMARY_HEADERS)]:
    ws = out_wb[sheet_name]
    if ws.max_row:
        ws.delete_rows(1, ws.max_row)
    ws.append(header_row)

raw_ws = out_wb['RawData']
formatted_ws = out_wb['Formatted Data']
summary_ws = out_wb['Summary']

agg = defaultdict(lambda: [0, 0, 0])
shipment_totals = defaultdict(int)
total_missing = total_zone = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    missing_scan, zone_mismatch, total, summary = calc_flags(row, latest_loaded)
    formatted_ws.append(base_values + [missing_scan, zone_mismatch, total, summary])
    key = (str(row['Route']), str(row['Shipment ID']))
    agg[key][0] += missing_scan
    agg[key][1] += zone_mismatch
    agg[key][2] += total
    shipment_totals[str(row['Shipment ID'])] += total
    total_missing += missing_scan
    total_zone += zone_mismatch
    total_errors += total

for (route, shipment_id), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([route, shipment_id, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_missing, total_zone, total_errors])
out_wb.save(OUTPUT_XLSX)

top_shipments = [sid for sid, total in sorted(shipment_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('Missing Load Scan checks whether each planned carton has a retained LOADED scan in the dock log. Zone Mismatch checks whether the retained LOADED scan points to a different zone than the planned outbound zone.')
doc.add_paragraph(f'The audit found {total_missing} Missing Load Scans, {total_zone} Zone Mismatches, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority shipment IDs include {", ".join(top_shipments)}. Recommendation: review the repeated route exceptions first and tighten dock controls so each carton gets a final LOADED scan in the planned zone.')
doc.save(OUTPUT_DOCX)
