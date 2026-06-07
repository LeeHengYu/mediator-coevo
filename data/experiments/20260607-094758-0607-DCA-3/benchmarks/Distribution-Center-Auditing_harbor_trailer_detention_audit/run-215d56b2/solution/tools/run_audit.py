#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
SOURCE_XLSX = ROOT_DIR / 'Trailer_Detention_Log.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Trailer_Detention_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Trailer_Detention_Brief.docx'

BASE_HEADERS = [
    'Load ID',
    'Carrier',
    'Allowed Hold Hours',
    'Actual Hold Hours',
    'Seal Required',
    'Seal Status',
    'Yard',
    'Dispatcher',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Detention Overrun', 'Seal Error', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Carrier', 'Yard', 'Detention Overrun Errors', 'Seal Errors', 'Total Errors']


def calc_flags(row):
    detention = 1 if float(row['Actual Hold Hours']) > float(row['Allowed Hold Hours']) else 0
    seal_required = str(row['Seal Required'] or '').strip().upper()
    seal_status = str(row['Seal Status'] or '').strip().upper()
    seal_error = 1 if seal_required == 'YES' and seal_status != 'VERIFIED' else 0
    total = detention + seal_error
    if total == 0:
        summary = 'None'
    elif detention and seal_error:
        summary = 'Detention Overrun, Seal Error'
    elif detention:
        summary = 'Detention Overrun'
    else:
        summary = 'Seal Error'
    return detention, seal_error, total, summary


wb = load_workbook(SOURCE_XLSX, data_only=True)
ws = wb['Detention']
headers = [ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, ws.max_row + 1):
    values = [ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

out_wb = Workbook()
raw_ws = out_wb.active
raw_ws.title = 'RawData'
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
carrier_totals = defaultdict(int)
total_detention = total_seal = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    detention, seal, total, summary = calc_flags(row)
    formatted_ws.append(base_values + [detention, seal, total, summary])
    key = (str(row['Carrier']), str(row['Yard']))
    agg[key][0] += detention
    agg[key][1] += seal
    agg[key][2] += total
    carrier_totals[str(row['Carrier'])] += total
    total_detention += detention
    total_seal += seal
    total_errors += total

for (carrier, yard), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([carrier, yard, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_detention, total_seal, total_errors])
out_wb.save(OUTPUT_XLSX)

top_carriers = [c for c, total in sorted(carrier_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('Detention Overrun checks whether the actual hold hours exceed the allowed threshold for a trailer. Seal Error checks whether a seal-required load does not have a VERIFIED seal status.')
doc.add_paragraph(f'The audit found {total_detention} Detention Overrun errors, {total_seal} Seal Errors, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority carriers include {", ".join(top_carriers)}. Recommendation: review detention billing first and tighten seal-verification procedures before release.')
doc.save(OUTPUT_DOCX)
