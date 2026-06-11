#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
PLAN_XLSX = ROOT_DIR / 'Cycle_Plan.xlsx'
EVENT_XLSX = ROOT_DIR / 'Count_Event_Log.xlsx'
TEMPLATE_XLSX = ROOT_DIR / 'Cycle_Template.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Cycle_Count_Variance_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Cycle_Count_Variance_Brief.docx'

BASE_HEADERS = [
    'Facility',
    'Session ID',
    'Bin ID',
    'Product ID',
    'Expected Qty',
    'Allowed Variance',
    'Approval Needed',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Missing Final Count', 'Approval Gap', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Facility', 'Session ID', 'Missing Final Counts', 'Approval Gaps', 'Total Errors']


def calc_flags(row, latest_final):
    key = (str(row['Facility']), str(row['Session ID']), str(row['Bin ID']))
    event = latest_final.get(key)
    missing = 1 if event is None else 0
    approval_needed = str(row['Approval Needed'] or '').strip().upper() == 'YES'
    gap = 0
    if event is not None and approval_needed:
        expected = float(row['Expected Qty'])
        actual = event[1]
        allowed = float(row['Allowed Variance'])
        if abs(expected - actual) > allowed:
            gap = 1
    total = missing + gap
    if total == 0:
        summary = 'None'
    elif missing and gap:
        summary = 'Missing Final Count, Approval Gap'
    elif missing:
        summary = 'Missing Final Count'
    else:
        summary = 'Approval Gap'
    return missing, gap, total, summary


plan_wb = load_workbook(PLAN_XLSX, data_only=True)
event_wb = load_workbook(EVENT_XLSX, data_only=True)
template_wb = load_workbook(TEMPLATE_XLSX, data_only=True)

plan_ws = plan_wb['PlanLines']
event_ws = event_wb['Events']
template_ws = template_wb['Overview']

headers = [plan_ws.cell(1, c).value for c in range(1, 8)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, plan_ws.max_row + 1):
    values = [plan_ws.cell(r, c).value for c in range(1, 8)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

latest_final = {}
for r in range(2, event_ws.max_row + 1):
    facility = event_ws.cell(r, 1).value
    session = event_ws.cell(r, 2).value
    bin_id = event_ws.cell(r, 3).value
    event_time = event_ws.cell(r, 4).value
    event_type = str(event_ws.cell(r, 5).value or '').strip().upper()
    count_qty = event_ws.cell(r, 6).value
    if facility is None or session is None or bin_id is None or count_qty is None:
        continue
    if event_type != 'FINAL':
        continue
    key = (str(facility), str(session), str(bin_id))
    record = (str(event_time), float(count_qty))
    if key not in latest_final or record[0] > latest_final[key][0]:
        latest_final[key] = record

out_wb = Workbook()
overview_ws = out_wb.active
overview_ws.title = 'Overview'
for r in range(1, template_ws.max_row + 1):
    for c in range(1, template_ws.max_column + 1):
        overview_ws.cell(r, c).value = template_ws.cell(r, c).value

raw_ws = out_wb.create_sheet('RawData')
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
session_totals = defaultdict(int)
total_missing = total_gap = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    missing, gap, total, summary = calc_flags(row, latest_final)
    formatted_ws.append(base_values + [missing, gap, total, summary])
    key = (str(row['Facility']), str(row['Session ID']))
    agg[key][0] += missing
    agg[key][1] += gap
    agg[key][2] += total
    session_totals[key] += total
    total_missing += missing
    total_gap += gap
    total_errors += total

for (facility, session), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([facility, session, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_missing, total_gap, total_errors])
out_wb.save(OUTPUT_XLSX)

top_sessions = [f"{fac}-{sess}" for (fac, sess), total in sorted(session_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('Missing Final Count checks whether a bin has no final count event. Approval Gap checks whether an approval-needed bin has a variance exceeding the allowed threshold.')
doc.add_paragraph(f'The audit found {total_missing} Missing Final Counts, {total_gap} Approval Gaps, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority facility-session combinations include {", ".join(top_sessions)}. Recommendation: prioritize recount for high-variance bins and enforce final-event completion before session close.')
doc.save(OUTPUT_DOCX)
