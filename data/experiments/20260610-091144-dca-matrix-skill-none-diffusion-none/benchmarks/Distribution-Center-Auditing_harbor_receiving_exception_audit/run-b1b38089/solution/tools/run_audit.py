#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
SOURCE_XLSX = ROOT_DIR / 'Receiving_Log.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Receiving_Exception_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Receiving_Exception_Brief.docx'

BASE_HEADERS = [
    'Receipt ID',
    'Item Code',
    'Expected Qty',
    'Received Qty',
    'Storage Class',
    'Temp Status',
    'Supplier',
    'Dock',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Qty Variance', 'Cold Chain Error', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Item Code', 'Supplier', 'Qty Variance Errors', 'Cold Chain Errors', 'Total Errors']

def calc_flags(row):
    qty_variance = 1 if row['Received Qty'] != row['Expected Qty'] else 0
    storage_class = str(row['Storage Class'] or '').strip().upper()
    temp_status = str(row['Temp Status'] or '').strip().upper()
    cold_chain = 1 if storage_class in {'CHILLED', 'FROZEN'} and temp_status != 'OK' else 0
    total = qty_variance + cold_chain
    if total == 0:
        summary = 'None'
    elif qty_variance and cold_chain:
        summary = 'Qty Variance, Cold Chain Error'
    elif qty_variance:
        summary = 'Qty Variance'
    else:
        summary = 'Cold Chain Error'
    return qty_variance, cold_chain, total, summary

wb = load_workbook(SOURCE_XLSX, data_only=True)
ws = wb['Receipts']
headers = [ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, ws.max_row + 1):
    values = [ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

out_wb = Workbook()
raw_ws = out_wb.active
raw_ws.title = 'RawData'
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
item_totals = defaultdict(int)
total_qty = total_cold = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    qty, cold, total, summary = calc_flags(row)
    formatted_ws.append(base_values + [qty, cold, total, summary])
    key = (str(row['Item Code']), str(row['Supplier']))
    agg[key][0] += qty
    agg[key][1] += cold
    agg[key][2] += total
    item_totals[str(row['Item Code'])] += total
    total_qty += qty
    total_cold += cold
    total_errors += total

for (item_code, supplier), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([item_code, supplier, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_qty, total_cold, total_errors])
out_wb.save(OUTPUT_XLSX)

top_items = [
    item for item, total in sorted(item_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0
][:3]
doc = Document()
doc.add_paragraph('Qty Variance checks whether the received quantity differs from the expected quantity on a receipt line. Cold Chain Error checks chilled or frozen lines where the temperature status is not OK.')
doc.add_paragraph(f'The audit found {total_qty} Qty Variance errors, {total_cold} Cold Chain Error exceptions, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority item codes include {", ".join(top_items)}. Recommendation: review recurring dock exceptions first and tighten receiving checks for refrigerated loads before closeout.')
doc.save(OUTPUT_DOCX)
