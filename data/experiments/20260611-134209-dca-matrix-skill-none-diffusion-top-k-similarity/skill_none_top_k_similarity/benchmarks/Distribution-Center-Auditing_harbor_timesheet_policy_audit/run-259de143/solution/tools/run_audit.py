#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
SOURCE_XLSX = ROOT_DIR / 'Timesheet_Submissions.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Timesheet_Compliance_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Timesheet_Compliance_Brief.docx'

BASE_HEADERS = [
    'Week Ending',
    'Employee ID',
    'Role',
    'Hours Worked',
    'Break Minutes',
    'Approval Code',
    'Project Code',
    'Manager',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Break Deficit', 'Approval Missing', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Employee ID', 'Week Ending', 'Break Deficits', 'Approval Gaps', 'Total Errors']

def calc_flags(row, rules):
    rule = rules[str(row['Role'])]
    break_deficit = 1 if float(row['Break Minutes']) < rule['min_break'] else 0
    approval_code = str(row['Approval Code'] or '').strip()
    approval_missing = 1 if float(row['Hours Worked']) > rule['overtime_threshold'] and not approval_code else 0
    total = break_deficit + approval_missing
    if total == 0:
        summary = 'None'
    elif break_deficit and approval_missing:
        summary = 'Break Deficit, Approval Missing'
    elif break_deficit:
        summary = 'Break Deficit'
    else:
        summary = 'Approval Missing'
    return break_deficit, approval_missing, total, summary

wb = load_workbook(SOURCE_XLSX, data_only=True)
entries_ws = wb['Entries']
rules_ws = wb['BreakRules']
headers = [entries_ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, entries_ws.max_row + 1):
    values = [entries_ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

rules = {}
for r in range(2, rules_ws.max_row + 1):
    role = rules_ws.cell(r, 1).value
    if role is None:
        continue
    rules[str(role)] = {
        'min_break': float(rules_ws.cell(r, 2).value),
        'overtime_threshold': float(rules_ws.cell(r, 3).value),
    }

out_wb = Workbook()
raw_ws = out_wb.active
raw_ws.title = 'RawData'
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
employee_totals = defaultdict(int)
total_break = total_approval = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    break_deficit, approval_missing, total, summary = calc_flags(row, rules)
    formatted_ws.append(base_values + [break_deficit, approval_missing, total, summary])
    key = (str(row['Employee ID']), str(row['Week Ending']))
    agg[key][0] += break_deficit
    agg[key][1] += approval_missing
    agg[key][2] += total
    employee_totals[str(row['Employee ID'])] += total
    total_break += break_deficit
    total_approval += approval_missing
    total_errors += total

for (employee_id, week_ending), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([employee_id, week_ending, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_break, total_approval, total_errors])
out_wb.save(OUTPUT_XLSX)

top_employees = [emp for emp, total in sorted(employee_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('Break Deficit checks whether the submitted break minutes fall below the minimum break requirement for the employee role. Approval Missing checks overtime entries that exceed the role threshold without an approval code.')
doc.add_paragraph(f'The audit found {total_break} Break Deficits, {total_approval} Approval Gaps, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority employee IDs include {", ".join(top_employees)}. Recommendation: review repeated break shortfalls first and require managers to clear overtime approvals before payroll lock.')
doc.save(OUTPUT_DOCX)
