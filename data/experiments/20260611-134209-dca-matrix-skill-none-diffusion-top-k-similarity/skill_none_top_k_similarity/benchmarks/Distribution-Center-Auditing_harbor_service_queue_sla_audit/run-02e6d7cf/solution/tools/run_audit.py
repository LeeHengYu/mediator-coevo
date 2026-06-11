#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
SOURCE_XLSX = ROOT_DIR / 'Ticket_Queue.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Service_Queue_SLA_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Service_Queue_SLA_Brief.docx'

BASE_HEADERS = [
    'Ticket ID',
    'Queue',
    'Priority Tier',
    'Open Age Hours',
    'Owner',
    'Escalation Code',
    'Region',
    'Analyst',
]
FORMATTED_HEADERS = BASE_HEADERS + ['SLA Breach', 'Missing Escalation', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['Queue', 'Region', 'SLA Breaches', 'Missing Escalations', 'Total Errors']


def calc_flags(row, rules):
    rule = rules[str(row['Priority Tier'])]
    sla_breach = 1 if float(row['Open Age Hours']) > rule['max_hours'] else 0
    escalation_code = str(row['Escalation Code'] or '').strip()
    missing_escalation = 1 if rule['escalation_required'] and not escalation_code else 0
    total = sla_breach + missing_escalation
    if total == 0:
        summary = 'None'
    elif sla_breach and missing_escalation:
        summary = 'SLA Breach, Missing Escalation'
    elif sla_breach:
        summary = 'SLA Breach'
    else:
        summary = 'Missing Escalation'
    return sla_breach, missing_escalation, total, summary


wb = load_workbook(SOURCE_XLSX, data_only=True)
tickets_ws = wb['Tickets']
rules_ws = wb['SLA_Rules']
headers = [tickets_ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, tickets_ws.max_row + 1):
    values = [tickets_ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

rules = {}
for r in range(2, rules_ws.max_row + 1):
    tier = rules_ws.cell(r, 1).value
    if tier is None:
        continue
    rules[str(tier)] = {
        'max_hours': float(rules_ws.cell(r, 2).value),
        'escalation_required': str(rules_ws.cell(r, 3).value or '').strip().upper() == 'Y',
    }

out_wb = Workbook()
raw_ws = out_wb.active
raw_ws.title = 'RawData'
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
queue_totals = defaultdict(int)
total_sla = total_missing = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    sla, missing, total, summary = calc_flags(row, rules)
    formatted_ws.append(base_values + [sla, missing, total, summary])
    key = (str(row['Queue']), str(row['Region']))
    agg[key][0] += sla
    agg[key][1] += missing
    agg[key][2] += total
    queue_totals[str(row['Queue'])] += total
    total_sla += sla
    total_missing += missing
    total_errors += total

for (queue, region), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([queue, region, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_sla, total_missing, total_errors])
out_wb.save(OUTPUT_XLSX)

top_queues = [q for q, total in sorted(queue_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('SLA Breach checks whether the ticket open age exceeds the maximum allowed hours for its priority tier. Missing Escalation checks whether a priority tier requires escalation but the ticket has no escalation code.')
doc.add_paragraph(f'The audit found {total_sla} SLA Breaches, {total_missing} Missing Escalations, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority queues include {", ".join(top_queues)}. Recommendation: review SLA thresholds for overdue tickets first and enforce escalation-code entry before SLA lockout.')
doc.save(OUTPUT_DOCX)
