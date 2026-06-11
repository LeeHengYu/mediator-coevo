#!/usr/bin/env python3
from collections import defaultdict
from pathlib import Path
import datetime as dt
import sys

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/root')
SOURCE_XLSX = ROOT_DIR / 'Promo_Price_Check_Source.xlsx'
OUTPUT_XLSX = ROOT_DIR / 'Promo_Register_Audit.xlsx'
OUTPUT_DOCX = ROOT_DIR / 'Promo_Register_Brief.docx'

BASE_HEADERS = [
    'Promo ID',
    'SKU',
    'Promo Price',
    'Register Price',
    'Promo Start Date',
    'Sale Date',
    'Promo End Date',
    'Store ID',
]
FORMATTED_HEADERS = BASE_HEADERS + ['Price Error', 'Window Error', 'Total Errors', 'Error Summary']
SUMMARY_HEADERS = ['SKU', 'Store ID', 'Price Errors', 'Window Errors', 'Total Errors']

def parse_date(value):
    if hasattr(value, 'date'):
        return value.date()
    if isinstance(value, dt.date):
        return value
    return dt.date.fromisoformat(str(value))

def calc_flags(row):
    price_error = 1 if row['Register Price'] != row['Promo Price'] else 0
    start_date = parse_date(row['Promo Start Date'])
    sale_date = parse_date(row['Sale Date'])
    end_date = parse_date(row['Promo End Date'])
    window_error = 1 if sale_date < start_date or sale_date > end_date else 0
    total = price_error + window_error
    if total == 0:
        summary = 'None'
    elif price_error and window_error:
        summary = 'Price Error, Window Error'
    elif price_error:
        summary = 'Price Error'
    else:
        summary = 'Window Error'
    return price_error, window_error, total, summary

wb = load_workbook(SOURCE_XLSX, data_only=True)
ws = wb['SalesAudit']
headers = [ws.cell(1, c).value for c in range(1, 9)]
if headers != BASE_HEADERS:
    raise ValueError(f'Unexpected source headers: {headers}')

rows = []
for r in range(2, ws.max_row + 1):
    values = [ws.cell(r, c).value for c in range(1, 9)]
    if all(v is None for v in values):
        continue
    rows.append(dict(zip(BASE_HEADERS, values)))

out_wb = Workbook()
raw_ws = out_wb.active
raw_ws.title = 'RawData'
raw_ws.append(BASE_HEADERS)
formatted_ws = out_wb.create_sheet('Formatted Data')
formatted_ws.append(FORMATTED_HEADERS)
summary_ws = out_wb.create_sheet('Summary')
summary_ws.append(SUMMARY_HEADERS)

agg = defaultdict(lambda: [0, 0, 0])
sku_totals = defaultdict(int)
total_price = total_window = total_errors = 0

for row in rows:
    base_values = [row[h] for h in BASE_HEADERS]
    raw_ws.append(base_values)
    price, window, total, summary = calc_flags(row)
    formatted_ws.append(base_values + [price, window, total, summary])
    key = (str(row['SKU']), str(row['Store ID']))
    agg[key][0] += price
    agg[key][1] += window
    agg[key][2] += total
    sku_totals[str(row['SKU'])] += total
    total_price += price
    total_window += window
    total_errors += total

for (sku, store_id), values in sorted(agg.items(), key=lambda x: (x[0][0], x[0][1])):
    if values[2] > 0:
        summary_ws.append([sku, store_id, values[0], values[1], values[2]])
summary_ws.append(['Grand Total', '-', total_price, total_window, total_errors])
out_wb.save(OUTPUT_XLSX)

top_skus = [sku for sku, total in sorted(sku_totals.items(), key=lambda x: (-x[1], x[0])) if total > 0][:3]
doc = Document()
doc.add_paragraph('Price Error checks whether the register price differs from the planned promo price. Window Error checks whether the sale date falls outside the approved promo start and end dates.')
doc.add_paragraph(f'The audit found {total_price} Price Errors, {total_window} Window Errors, and {total_errors} total errors.')
doc.add_paragraph(f'High-priority SKUs include {", ".join(top_skus)}. Recommendation: review the repeated register overrides first and tighten controls on selling outside approved promo windows.')
doc.save(OUTPUT_DOCX)
